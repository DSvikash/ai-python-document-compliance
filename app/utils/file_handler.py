"""
File handling utilities for document processing.
"""
import os
import uuid
from pathlib import Path
from typing import Tuple, Optional
import PyPDF2
import docx
from fastapi import UploadFile, HTTPException


class FileHandler:
    """Handles file upload, validation, and text extraction."""
    
    def __init__(self, upload_dir: Path, allowed_extensions: list, max_file_size: int):
        """
        Initialize FileHandler.
        
        Args:
            upload_dir: Directory to store uploaded files
            allowed_extensions: List of allowed file extensions
            max_file_size: Maximum file size in bytes
        """
        self.upload_dir = upload_dir
        self.allowed_extensions = allowed_extensions
        self.max_file_size = max_file_size
        self.upload_dir.mkdir(exist_ok=True)
    
    def validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file
            
        Raises:
            HTTPException: If file validation fails
        """
        # Check file extension
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in self.allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
            )
        
        # Check file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > self.max_file_size:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {self.max_file_size / (1024*1024):.2f}MB"
            )
    
    async def save_file(self, file: UploadFile) -> Tuple[str, Path]:
        """
        Save uploaded file to disk.
        
        Args:
            file: Uploaded file
            
        Returns:
            Tuple of (document_id, file_path)
        """
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        file_ext = file.filename.split('.')[-1].lower()
        file_path = self.upload_dir / f"{document_id}.{file_ext}"
        
        # Save file
        content = await file.read()
        with open(file_path, 'wb') as f:
            f.write(content)
        
        return document_id, file_path
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting text from PDF: {str(e)}"
            )
        return text
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting text from DOCX: {str(e)}"
            )
        return text
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_ext}"
            )
    
    def get_file_path(self, document_id: str) -> Optional[Path]:
        """
        Get file path for a document ID.
        
        Args:
            document_id: Document identifier
            
        Returns:
            File path if exists, None otherwise
        """
        for ext in self.allowed_extensions:
            file_path = self.upload_dir / f"{document_id}.{ext}"
            if file_path.exists():
                return file_path
        return None
    
    def create_modified_document(self, original_path: Path, modified_text: str) -> Path:
        """
        Create a modified document with updated text.
        
        Args:
            original_path: Path to original document
            modified_text: Modified text content
            
        Returns:
            Path to modified document
        """
        file_ext = original_path.suffix.lower()
        document_id = str(uuid.uuid4())
        modified_path = self.upload_dir / f"{document_id}_modified{file_ext}"
        
        if file_ext == '.docx':
            doc = docx.Document()
            for paragraph in modified_text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            doc.save(modified_path)
        else:
            # For PDF, we'll create a simple text file for now
            # In production, you'd use a library like reportlab
            modified_path = self.upload_dir / f"{document_id}_modified.txt"
            with open(modified_path, 'w', encoding='utf-8') as f:
                f.write(modified_text)
        
        return modified_path
